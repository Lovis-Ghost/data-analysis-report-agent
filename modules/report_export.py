import io

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def _is_markdown_table_separator(line):
    cleaned_line = line.replace("|", "").replace("-", "").replace(":", "").strip()
    return cleaned_line == ""


def _split_markdown_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _get_markdown_table_rows(table_lines):
    table_rows = []

    for line in table_lines:
        if _is_markdown_table_separator(line):
            continue

        row_cells = _split_markdown_table_row(line)
        if row_cells:
            table_rows.append(row_cells)

    return table_rows


def _add_markdown_table(document, table_lines):
    table_rows = _get_markdown_table_rows(table_lines)

    if not table_rows:
        return

    column_count = max(len(row) for row in table_rows)
    table = document.add_table(rows=len(table_rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row_cells in enumerate(table_rows):
        for col_index in range(column_count):
            cell_text = row_cells[col_index] if col_index < len(row_cells) else ""
            table.cell(row_index, col_index).text = cell_text


def create_docx_report_bytes(report_text):
    document = Document()
    lines = report_text.splitlines()
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            _add_markdown_table(document, table_lines)
            table_lines = []

    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith("|"):
            table_lines.append(stripped_line)
            continue

        flush_table()

        if not stripped_line:
            continue

        if stripped_line.startswith("# "):
            document.add_heading(stripped_line[2:].strip(), level=0)
        elif stripped_line.startswith("## "):
            document.add_heading(stripped_line[3:].strip(), level=1)
        elif stripped_line.startswith("### "):
            document.add_heading(stripped_line[4:].strip(), level=2)
        elif stripped_line.startswith("- "):
            document.add_paragraph(stripped_line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped_line)

    flush_table()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _clean_pdf_text(text):
    replacements = {
        "→": "->",
        "²": "2",
        "–": "-",
        "—": "-",
        "“": '"',
        "”": '"',
        "’": "'",
        "•": "-",
    }

    for old_value, new_value in replacements.items():
        text = text.replace(old_value, new_value)

    return text.encode("latin-1", errors="replace").decode("latin-1")


def _add_pdf_table(story, table_lines, styles):
    table_rows = _get_markdown_table_rows(table_lines)

    if not table_rows:
        return

    column_count = max(len(row) for row in table_rows)
    table_data = []

    for row in table_rows:
        padded_row = row + [""] * (column_count - len(row))
        table_data.append([
            Paragraph(_clean_pdf_text(cell), styles["TableCell"])
            for cell in padded_row
        ])

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.12 * inch))


def create_pdf_report_bytes(report_text):
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="BulletText",
        parent=styles["BodyText"],
        leftIndent=14,
        firstLineIndent=-8,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="TableCell",
        parent=styles["BodyText"],
        fontSize=8,
        leading=10,
    ))

    story = []
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            _add_pdf_table(story, table_lines, styles)
            table_lines = []

    for line in report_text.splitlines():
        stripped_line = line.strip()

        if stripped_line.startswith("|"):
            table_lines.append(stripped_line)
            continue

        flush_table()

        if not stripped_line:
            continue

        safe_line = _clean_pdf_text(stripped_line)

        if safe_line.startswith("# "):
            story.append(Paragraph(safe_line[2:].strip(), styles["Title"]))
            story.append(Spacer(1, 0.16 * inch))
        elif safe_line.startswith("## "):
            story.append(Paragraph(safe_line[3:].strip(), styles["Heading1"]))
        elif safe_line.startswith("### "):
            story.append(Paragraph(safe_line[4:].strip(), styles["Heading2"]))
        elif safe_line.startswith("- "):
            story.append(Paragraph(f"- {safe_line[2:].strip()}", styles["BulletText"]))
        else:
            story.append(Paragraph(safe_line, styles["BodyText"]))

    flush_table()

    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()
